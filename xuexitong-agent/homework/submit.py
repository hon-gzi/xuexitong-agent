import asyncio
import os
import time
from browser.driver import BrowserDriver
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HomeworkSubmitter:
    """作业提交"""

    def __init__(self, driver: BrowserDriver):
        self.driver = driver

    async def fill_and_submit(self, questions: list[dict], answers: list[str], frame=None) -> dict:
        """
        填写答案并提交

        questions: 题目列表（包含containerIndex信息）
        answers: 对应的答案列表
        frame: 题目所在的 frame（可选，不传则自动搜索）
        """
        if frame:
            target = frame
        else:
            # 自动搜索包含题目的 frame
            target = self.driver.page
            for f in self.driver.page.frames:
                try:
                    has_q = await f.evaluate("""
                        () => document.querySelectorAll('.questionLi, .singleQuesId, .mark_item, .Zy_TIt6').length
                    """)
                    if has_q > 0:
                        target = f
                        print(f"[提交] 找到题目frame: {f.url[:60]}")
                        break
                except Exception:
                    continue
        filled = 0

        # 用 JS 一次性填写所有答案（比逐个 locator 更可靠）
        fill_data = []
        for q, answer in zip(questions, answers):
            idx = q.get("containerIndex", q.get("dom_index", q.get("index", -1)))
            fill_data.append({
                "index": idx,
                "type": q.get("type", "unknown"),
                "answer": answer,
            })

        try:
            result = await target.evaluate("""
                (data) => {
                    const selectors = '.questionLi, .singleQuesId, .mark_item, .Zy_TIt6';
                    let containers = document.querySelectorAll(selectors);
                    containers = Array.from(containers).filter(el => !el.classList.contains('fanyaMarking'));
                    let filled = 0;
                    let errors = [];

                    // 学习通选项结构: div[role="radio"] 内含 span[data="A"] 和 div.answer_p
                    // 点击 div[role="radio"] 即可选中（触发 onclick="addChoice(this)"）
                    function findOptionEls(c) {
                        // 1. 学习通标准结构：div[role="radio"]
                        const radios = c.querySelectorAll('[role="radio"]');
                        if (radios.length >= 2) return Array.from(radios);
                        // 2. 兜底：选项专用 class
                        const specific = c.querySelectorAll('.xuanxiang, .option, [data-key]');
                        if (specific.length >= 2) return Array.from(specific);
                        // 3. 最后兜底：li
                        return Array.from(c.querySelectorAll('li'));
                    }

                    // 根据答案字母点击对应选项
                    function clickByLetter(optionEls, letter) {
                        const L = letter.toUpperCase();
                        // 方式1: 找 span[data="A"] 所在的父级 radio div
                        for (const el of optionEls) {
                            const span = el.querySelector('span[data]');
                            if (span && span.getAttribute('data').toUpperCase() === L) {
                                el.click(); return true;
                            }
                        }
                        // 方式2: 精确文本匹配（文本以 "A." 开头）
                        for (const el of optionEls) {
                            const t = el.innerText.trim();
                            if (new RegExp('^' + L + '[.、．\\\\s]').test(t)) {
                                el.click(); return true;
                            }
                        }
                        // 方式3: 按位置索引（A=0, B=1, ...）
                        const i = L.charCodeAt(0) - 65;
                        if (i >= 0 && i < optionEls.length) {
                            optionEls[i].click(); return true;
                        }
                        return false;
                    }

                    // 判断题：span[data="true"] 对应"对"，span[data="false"] 对应"错"
                    function clickJudge(c, answer) {
                        const wantTrue = answer.includes('对');
                        const radios = c.querySelectorAll('[role="radio"]');
                        for (const r of radios) {
                            const span = r.querySelector('span[data]');
                            const d = span ? span.getAttribute('data') : null;
                            if ((wantTrue && d === 'true') || (!wantTrue && d === 'false')) {
                                r.click(); return true;
                            }
                        }
                        // 兜底：文本匹配
                        const btns = c.querySelectorAll('[role="radio"], li, span');
                        for (const b of btns) {
                            const t = b.innerText.trim();
                            if ((wantTrue && (t === '对' || t === '正确' || t === 'A')) ||
                                (!wantTrue && (t === '错' || t === '错误' || t === 'B'))) {
                                b.click(); return true;
                            }
                        }
                        return false;
                    }

                    for (const item of data) {
                        const idx = item.index;
                        if (idx < 0 || idx >= containers.length) {
                            errors.push('idx=' + idx + ' out of range');
                            continue;
                        }
                        const c = containers[idx];

                        try {
                            if (item.type === 'single') {
                                if (clickByLetter(findOptionEls(c), item.answer)) filled++;
                                else errors.push('single idx=' + idx + ' letter=' + item.answer);
                            }
                            else if (item.type === 'multi') {
                                const letters = item.answer.split(',').map(s => s.trim().toUpperCase());
                                for (const letter of letters) clickByLetter(findOptionEls(c), letter);
                                filled++;
                            }
                            else if (item.type === 'judge') {
                                if (clickJudge(c, item.answer)) filled++;
                                else errors.push('judge idx=' + idx + ' answer=' + item.answer);
                            }
                            else {
                                // 填空/简答/编程/论文
                                const input = c.querySelector('textarea, input[type="text"], [contenteditable="true"], .ueditor');
                                if (input) {
                                    if (input.tagName === 'TEXTAREA' || input.tagName === 'INPUT') {
                                        input.value = item.answer;
                                    } else {
                                        input.innerText = item.answer;
                                    }
                                    input.dispatchEvent(new Event('input', {bubbles: true}));
                                    input.dispatchEvent(new Event('change', {bubbles: true}));
                                    filled++;
                                } else {
                                    errors.push('no input for idx=' + idx);
                                }
                            }
                        } catch(e) {
                            errors.push('idx=' + idx + ' error: ' + e.message);
                        }
                    }
                    return {filled: filled, total: data.length, errors: errors};
                }
            """, fill_data)
            print(f"[提交] JS填写结果: {result}")
            filled = result.get("filled", 0)
        except Exception as e:
            print(f"[提交] JS填写失败: {e}")

        print(f"[填写] 已填写 {filled}/{len(questions)} 题")

        # 处理编程题和论文：保存到文件
        code_files = []
        paper_files = []
        for q, answer in zip(questions, answers):
            if q.get("type") == "code":
                filename = self._save_code_file(q, answer)
                if filename:
                    code_files.append(filename)
            elif q.get("type") == "paper":
                filename = self._save_paper_file(q, answer)
                if filename:
                    paper_files.append(filename)

        if code_files:
            print(f"\n[编程题] 已生成 {len(code_files)} 个代码文件：")
            for f in code_files:
                print(f"  - {f}")
            print(f"[编程题] 请手动上传这些文件到对应题目")

        if paper_files:
            print(f"\n[论文] 已生成 {len(paper_files)} 个论文文件：")
            for f in paper_files:
                print(f"  - {f}")
            print(f"[论文] 请手动上传这些文件到对应题目")

        print(f"[填写] 请手动检查并提交作业")

        return {"status": "filled_only", "filled": filled, "code_files": code_files, "paper_files": paper_files}

    async def _fill_question(self, page, question: dict, answer: str):
        """填写单个题目的答案"""
        q_type = question.get("type", "unknown")
        selector = question.get("selector", "")

        if q_type == "single":
            # 单选题：点击对应选项
            await self._click_option(page, selector, answer)

        elif q_type == "multi":
            # 多选题：点击多个选项
            for letter in answer.split(","):
                await self._click_option(page, selector, letter.strip())
                await self.driver.random_delay(0.3, 0.5)

        elif q_type == "judge":
            # 判断题：点击对/错
            await self._click_judge(page, selector, answer)

        elif q_type in ("fill", "short", "code"):
            # 填空/简答/编程：填写文本
            await self._fill_text(page, selector, answer)

    async def _click_option(self, page, q_selector: str, letter: str):
        """点击选择题选项"""
        # 尝试多种选择器
        selectors = [
            f"{q_selector} li:nth-child({ord(letter.upper()) - 64})",
            f"{q_selector} .xuanxiang:nth-child({ord(letter.upper()) - 64})",
            f"{q_selector} [data-key='{letter.upper()}']",
        ]
        for sel in selectors:
            try:
                el = page.locator(sel).first
                if await el.is_visible(timeout=1000):
                    await el.click()
                    return
            except Exception:
                continue

    async def _click_judge(self, page, q_selector: str, answer: str):
        """点击判断题"""
        if "对" in answer or "正确" in answer or answer.upper() == "A":
            selectors = [f"{q_selector} text=对", f"{q_selector} text=正确", f"{q_selector} text=A"]
        else:
            selectors = [f"{q_selector} text=错", f"{q_selector} text=错误", f"{q_selector} text=B"]

        for sel in selectors:
            try:
                el = page.locator(sel).first
                if await el.is_visible(timeout=1000):
                    await el.click()
                    return
            except Exception:
                continue

    async def _fill_text(self, page, q_selector: str, answer: str):
        """填写文本类答案"""
        # 查找输入框/文本区域
        selectors = [
            f"{q_selector} textarea",
            f"{q_selector} input[type='text']",
            f"{q_selector} .ueditor",
            f"{q_selector} [contenteditable='true']",
        ]
        for sel in selectors:
            try:
                el = page.locator(sel).first
                if await el.is_visible(timeout=1000):
                    await el.fill(answer)
                    return
            except Exception:
                continue

    def _save_code_file(self, question: dict, code: str) -> str:
        """保存编程题代码到文件"""
        # 创建代码目录
        code_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "code")
        os.makedirs(code_dir, exist_ok=True)

        # 生成文件名
        q_index = question.get("index", 0)
        q_text = question.get("text", "")[:30]
        # 清理文件名中的非法字符（保留中文、字母、数字）
        q_text = "".join(c for c in q_text if c.isalnum() or c in " _-" or '一' <= c <= '鿿').strip()
        # 如果清理后为空，使用默认名
        if not q_text:
            q_text = "code"
        timestamp = int(time.time())

        # 根据题目判断语言
        lang = "py"  # 默认 Python
        if "java" in q_text.lower():
            lang = "java"
        elif "c++" in q_text.lower() or "cpp" in q_text.lower():
            lang = "cpp"
        elif "c语言" in q_text.lower() or "c语言" in q_text:
            lang = "c"
        elif "javascript" in q_text.lower() or "js" in q_text.lower():
            lang = "js"

        filename = f"q{q_index}_{q_text}_{timestamp}.{lang}"
        filepath = os.path.join(code_dir, filename)

        # 保存代码
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(code)

        print(f"[编程题] 保存代码到: {filepath}")
        return filepath

    def _save_paper_file(self, question: dict, content: str) -> str:
        """保存论文作业到 Word 文件"""
        # 创建代码目录
        code_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "code")
        os.makedirs(code_dir, exist_ok=True)

        # 生成文件名
        q_index = question.get("index", 0)
        q_text = question.get("text", "")[:30]
        q_text = "".join(c for c in q_text if c.isalnum() or c in " _-" or '一' <= c <= '鿿').strip()
        if not q_text:
            q_text = "论文"
        timestamp = int(time.time())

        filename = f"q{q_index}_{q_text}_{timestamp}.docx"
        filepath = os.path.join(code_dir, filename)

        # 创建 Word 文档
        doc = Document()

        # 设置标题
        title = doc.add_heading(q_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加正文
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                # 检查是否是标题（以数字开头或包含"摘要"、"关键词"等）
                if any(kw in para_text for kw in ['摘要', '关键词', '引言', '结论', '参考文献']):
                    doc.add_heading(para_text, level=1)
                elif para_text.strip().startswith(('1.', '2.', '3.', '4.', '5.', '一、', '二、', '三、')):
                    doc.add_heading(para_text, level=2)
                else:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进

        # 保存文档
        doc.save(filepath)
        print(f"[论文] 保存论文到: {filepath}")
        return filepath
